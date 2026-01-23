import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement, ns
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime
import io
import re
from PyPDF2 import PdfReader
import mimetypes

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="Weekly Report Update", layout="centered")
st.title("📊 Weekly Report Update")

# ---------------- INPUTS ----------------
current_message = st.text_input("Current Message")
tracking = st.text_area("What are we tracking", height=150)
record_results = st.text_input("Where to record the results")
users = st.text_input("Who should use the records")
helpful = st.text_area("How this information would be helpful", height=120)
contacts = st.text_area("Important contacts", height=120)

attachments = st.file_uploader(
    "Attachments (Any format – multiple allowed)",
    accept_multiple_files=True
)

# ---------------- WORD HELPERS ----------------
def add_page_border(doc):
    section = doc.sections[0]
    sectPr = section._sectPr
    borders = OxmlElement('w:pgBorders')
    borders.set(ns.qn('w:offsetFrom'), 'page')

    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(ns.qn('w:val'), 'single')
        b.set(ns.qn('w:sz'), '12')
        b.set(ns.qn('w:color'), '000000')
        borders.append(b)

    sectPr.append(borders)


def add_title_with_border(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = 1

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(ns.qn('w:val'), 'single')
        b.set(ns.qn('w:sz'), '18')
        b.set(ns.qn('w:color'), '003366')
        pBdr.append(b)

    pPr.append(pBdr)


def extract_urls(text):
    return re.findall(r'(https?://\S+)', text)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(ns.qn('w:val'), '0000FF')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(ns.qn('w:val'), 'single')
    rPr.append(underline)

    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(doc, title, content):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()

    urls = extract_urls(content)
    modified_text = content

    for i, url in enumerate(urls, start=1):
        modified_text = modified_text.replace(url, f"[Link-{i}]")

    p.add_run(modified_text)

    for i, url in enumerate(urls, start=1):
        add_hyperlink(p, url, f" Link-{i} ")

# ---------------- ATTACHMENT HANDLER ----------------
def add_attachment_to_doc(doc, file):
    name = file.name
    mime_type, _ = mimetypes.guess_type(name)

    doc.add_page_break()
    doc.add_heading(f"📎 Attachment: {name}", level=2)

    # IMAGE FILES
    if mime_type and mime_type.startswith("image"):
        doc.add_picture(file, width=Inches(4))
        doc.add_paragraph(f"Image file: {name}")
        return

    # PDF FILES
    if name.lower().endswith(".pdf"):
        reader = PdfReader(file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
        return

    # WORD FILES
    if name.lower().endswith(".docx"):
        from docx import Document as ReadDoc
        temp_doc = ReadDoc(file)
        for p in temp_doc.paragraphs:
            doc.add_paragraph(p.text)
        return

    # TEXT-BASED FILES
    if name.lower().endswith(('.txt', '.csv', '.json', '.xml', '.log', '.md')):
        text = file.read().decode("utf-8", errors="ignore")
        doc.add_paragraph(text)
        return

    # BINARY FALLBACK (SAFE)
    raw = file.read()
    preview = raw[:2000].decode("utf-8", errors="ignore").replace("\x00", "")
    doc.add_paragraph(
        f"⚠ Binary file detected\nFilename: {name}\n\nPreview:\n{preview}"
    )

# ---------------- SUBMIT ----------------
if st.button("📥 Submit & Download Word Document"):
    doc = Document()

    add_page_border(doc)
    add_title_with_border(doc, "Weekly Report Update")

    add_text_with_links(doc, "Current Message", current_message)
    add_text_with_links(doc, "What are we tracking", tracking)
    add_text_with_links(doc, "Where to record the results", record_results)
    add_text_with_links(doc, "Who should use the records", users)
    add_text_with_links(doc, "How this information would be helpful", helpful)
    add_text_with_links(doc, "Important contacts", contacts)

    if attachments:
        for file in attachments:
            add_attachment_to_doc(doc, file)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Weekly_ReportUpdate_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    st.success("✅ Document generated successfully")

    st.download_button(
        "⬇ Download Word Document",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
